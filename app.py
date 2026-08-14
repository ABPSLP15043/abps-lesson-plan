import io
import json
import os
from xml.sax.saxutils import escape as xml_escape
import streamlit as st
import docx
import graphviz
from pypdf import PdfReader
from google import genai
from google.genai import types

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# ---------------------------------------------------------
# 1. PAGE CONFIG (Must be the very first Streamlit command)
# ---------------------------------------------------------
st.set_page_config(
    page_title="ABPS Baikunth - AI Lesson Plan Generator",
    page_icon="🏫",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------------------------------------
# 2. LANDING PAGE & UI STYLING
# ---------------------------------------------------------
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

html, body, [class*="css"]  {
    font-family: 'Inter', sans-serif !important;
}

.stApp {
    background:
        radial-gradient(circle at top left, rgba(59,130,246,0.18), transparent 30%),
        radial-gradient(circle at top right, rgba(168,85,247,0.16), transparent 28%),
        radial-gradient(circle at bottom left, rgba(16,185,129,0.14), transparent 26%),
        linear-gradient(135deg, #f8fbff 0%, #eef4ff 45%, #f7f5ff 100%);
    background-attachment: fixed;
}

.block-container {
    padding-top: 2rem;
    padding-bottom: 2rem;
}

h1 {
    background: linear-gradient(90deg, #0f172a 0%, #2563eb 45%, #7c3aed 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    font-weight: 800 !important;
    letter-spacing: -0.03em;
    margin-bottom: 0.2rem !important;
}

[data-testid="stCaptionContainer"] {
    color: #475569 !important;
    font-weight: 600 !important;
    font-size: 1rem !important;
}

[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0f172a 0%, #111827 100%) !important;
    border-right: 1px solid rgba(255,255,255,0.08);
}

[data-testid="stSidebar"] label,
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] div {
    color: #e5eefc !important;
}

[data-testid="stSidebar"] input,
[data-testid="stSidebar"] textarea,
[data-testid="stSidebar"] div[role="combobox"] {
    background: #1e293b !important;
    color: #f8fafc !important;
    border: 1px solid #334155 !important;
    border-radius: 12px !important;
}

div[data-baseweb="popover"] ul {
    background: #1e293b !important;
    border: 1px solid #334155 !important;
    border-radius: 12px !important;
}

div[data-baseweb="popover"] li {
    color: #f8fafc !important;
}

div[data-baseweb="popover"] li:hover {
    background: #334155 !important;
}

[data-testid="stTabs"] {
    background: rgba(255,255,255,0.62) !important;
    backdrop-filter: blur(18px) !important;
    -webkit-backdrop-filter: blur(18px) !important;
    border: 1px solid rgba(255,255,255,0.7) !important;
    border-radius: 22px !important;
    padding: 1.2rem !important;
    box-shadow: 0 16px 40px rgba(15,23,42,0.08) !important;
}

button[data-baseweb="tab"] {
    border-radius: 999px !important;
    color: #64748b !important;
    font-weight: 700 !important;
    transition: all 0.2s ease !important;
}

button[data-baseweb="tab"][aria-selected="true"] {
    background: linear-gradient(90deg, #2563eb, #7c3aed) !important;
    color: white !important;
    box-shadow: 0 8px 20px rgba(99,102,241,0.28) !important;
}

.stTextArea textarea, .stFileUploader {
    background: rgba(255,255,255,0.84) !important;
    border: 1px solid #dbe7ff !important;
    border-radius: 16px !important;
    color: #0f172a !important;
    box-shadow: 0 10px 24px rgba(15,23,42,0.04) !important;
}

div.stButton > button {
    border: none !important;
    border-radius: 16px !important;
    font-weight: 800 !important;
    padding: 0.75rem 1.2rem !important;
    background: linear-gradient(135deg, #2563eb 0%, #7c3aed 100%) !important;
    color: #ffffff !important;
    box-shadow: 0 14px 30px rgba(99,102,241,0.28) !important;
    transition: all 0.25s ease !important;
}

div.stButton > button:hover {
    transform: translateY(-2px) scale(1.01);
    box-shadow: 0 18px 34px rgba(99,102,241,0.38) !important;
}

div[data-testid="stDownloadButton"] > button {
    border-radius: 14px !important;
    font-weight: 700 !important;
    background: linear-gradient(135deg, #ecfdf5 0%, #d1fae5 100%) !important;
    color: #065f46 !important;
    border: 1px solid #a7f3d0 !important;
}

div[data-testid="stDownloadButton"] > button:hover {
    background: linear-gradient(135deg, #10b981 0%, #059669 100%) !important;
    color: white !important;
}
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# 3. CONFIG & MODEL PREFERENCES
# ---------------------------------------------------------
PREFERRED_MODELS = [
    "gemini-3.6-flash",
    "gemini-3.5-flash",
    "gemini-3.5-flash-lite",
    "gemini-3.1-flash-lite",
    "gemini-2.5-flash-lite",
    "gemini-2.5-flash",
]

_STR_LIST = {"type": "array", "items": {"type": "string"}}

LESSON_PLAN_SCHEMA = {
    "type": "object",
    "properties": {
        "curriculum_goal": {"type": "string"},
        "relevant_competencies": _STR_LIST,
        "learning_objectives": _STR_LIST,
        "expected_learning_outcomes": _STR_LIST,
        "formulas_and_equations": _STR_LIST,
        "teaching_methodology": _STR_LIST,
        "teaching_aids": _STR_LIST,
        "art_integration": _STR_LIST,
        "previous_knowledge": _STR_LIST,
        "innovative_techniques": _STR_LIST,
        "content_points": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "section": {"type": "string"},
                    "topics": _STR_LIST
                },
                "required": ["section", "topics"]
            }
        },
        "projects_experiential": _STR_LIST,
        "skills_acquired": _STR_LIST,
        "values_inculcated": _STR_LIST,
        "multiple_assessment": {
            "type": "object",
            "properties": {
                "oral_questions": _STR_LIST,
                "worksheet": _STR_LIST,
                "practical": _STR_LIST,
                "exit_ticket": _STR_LIST
            },
            "required": ["oral_questions", "worksheet", "practical", "exit_ticket"]
        },
        "class_work": _STR_LIST,
        "home_work": _STR_LIST,
        "remedial_measures": {
            "type": "object",
            "properties": {
                "slow_learners": _STR_LIST,
                "advanced_learners": _STR_LIST
            },
            "required": ["slow_learners", "advanced_learners"]
        },
        "resources": {
            "type": "object",
            "properties": {
                "books": _STR_LIST,
                "websites": _STR_LIST,
                "videos": _STR_LIST
            },
            "required": ["books", "websites", "videos"]
        }
    },
    "required": [
        "curriculum_goal",
        "relevant_competencies",
        "learning_objectives",
        "expected_learning_outcomes",
        "formulas_and_equations",
        "teaching_methodology",
        "teaching_aids",
        "art_integration",
        "previous_knowledge",
        "innovative_techniques",
        "content_points",
        "projects_experiential",
        "skills_acquired",
        "values_inculcated",
        "multiple_assessment",
        "class_work",
        "home_work",
        "remedial_measures",
        "resources"
    ]
}

# ---------------------------------------------------------
# APPLICATION HEADER
# ---------------------------------------------------------
st.title("🏫 The Aditya Birla Public School, Baikunth")
st.caption("AI-Powered NCF-SE 2023 Lesson Plan & Mind Map Generator")

system_api_key = os.getenv("GEMINI_API_KEY", "")

# ---------------------------------------------------------
# SIDEBAR
# ---------------------------------------------------------
st.sidebar.header("📋 Lesson Plan Details")

teacher_name = st.sidebar.text_input("Teacher Name", "Educator Name")

user_api_key = st.sidebar.text_input(
    "Manual API Key (Optional)", type="password",
    help="Leave blank if GEMINI_API_KEY is set in your environment.",
)
active_api_key = user_api_key.strip() if user_api_key.strip() else system_api_key

subject = st.sidebar.selectbox("Subject", [
    "SCIENCE", "MATHEMATICS", "SOCIAL SCIENCE", "ENGLISH", "HINDI", "SANSKRIT",
    "MUSIC", "ART & CRAFT", "DANCE",
    "PHYSICS", "CHEMISTRY", "BIOLOGY", "COMPUTER SCIENCE", "IP",
    "ACCOUNTANCY", "BUSINESS STUDIES", "ECONOMICS", "HISTORY",
    "POLITICAL SCIENCE", "GEOGRAPHY",
    "ARTIFICIAL INTELLIGENCE (AI)", "INFORMATION TECHNOLOGY (IT)",
])

col_g, col_s = st.sidebar.columns(2)
with col_g:
    grade = st.selectbox("Class", ["VI", "VII", "VIII", "IX", "X", "XI", "XII"])
with col_s:
    section = st.text_input("Section", "A")

month = st.sidebar.selectbox("Month", [
    "APRIL", "MAY", "JULY", "AUGUST", "SEPTEMBER", "OCTOBER",
    "NOVEMBER", "DECEMBER", "JANUARY", "FEBRUARY",
])
chapter = st.sidebar.text_input("Chapter / Topic", "Force and Laws of Motion")
periods = st.sidebar.number_input("No. of Periods", min_value=1, max_value=25, value=8)

st.sidebar.divider()

if st.sidebar.button("🔍 Check available models"):
    if not active_api_key:
        st.sidebar.error("Enter an API key first.")
    else:
        try:
            diag_client = genai.Client(api_key=active_api_key)
            names = []
            for m in diag_client.models.list():
                raw = getattr(m, "name", str(m))
                names.append(raw.replace("models/", ""))
            if names:
                st.sidebar.success(f"{len(names)} models available:")
                st.sidebar.code("\n".join(sorted(names)))
            else:
                st.sidebar.warning("No models returned for this key.")
        except Exception as exc:
            st.sidebar.error(f"Could not list models: {exc}")

# ---------------------------------------------------------
# CHAPTER INPUT
# ---------------------------------------------------------
st.subheader("📂 Step 1: Provide Chapter Content")
tab1, tab2 = st.tabs(["📄 Upload Chapter PDF (Max 25MB)", "📝 Paste Chapter Text/Notes"])

uploaded_pdf = None
pasted_text = ""
with tab1:
    uploaded_pdf = st.file_uploader(
        "Upload NCERT / Textbook Chapter PDF (up to 25 MB)", type=["pdf"]
    )
with tab2:
    pasted_text = st.text_area(
        "Paste text, outline, or key topics from the chapter here:", height=200
    )

# ---------------------------------------------------------
# HELPERS & MODEL RESOLUTION
# ---------------------------------------------------------
@st.cache_data(show_spinner=False)
def extract_text_cached(file_bytes):
    reader = PdfReader(io.BytesIO(file_bytes))
    parts = []
    for page in reader.pages[:20]:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n".join(parts)

def get_available_model(client):
    available = []
    for m in client.models.list():
        name = getattr(m, "name", "")
        if name.startswith("models/"):
            name = name.replace("models/", "")
        available.append(name)

    for preferred in PREFERRED_MODELS:
        if preferred in available:
            return preferred, available

    for name in available:
        if "flash" in name.lower():
            return name, available

    raise RuntimeError(
        "No supported Flash model found for this API key.\n\nAvailable models:\n- " +
        "\n- ".join(available[:50])
    )

def generate_ai_lesson_plan(api_key, subject, grade, section, chapter, month, periods, chapter_content):
    client = genai.Client(api_key=api_key)
    model_name, available_models = get_available_model(client)

    prompt = f"""
You are an expert curriculum designer for Indian schools following NCF-SE 2023 and NEP 2020.

Create a specific, chapter-based lesson plan.

Subject: {subject}
Class/Section: {grade} - {section}
Chapter: {chapter}
Month: {month}
Periods: {periods}

CHAPTER CONTENT:{chapter_content[:7000]}

Rules:
- Be specific to the chapter
- Avoid generic placeholders
- Keep outputs concise and practical for teachers
- If no formulas exist, return an empty list
- Keep most lists between 3 and 6 items
"""

    if hasattr(client, "interactions"):
        result = client.interactions.create(
            model=model_name,
            input=prompt,
            response_format={
                "type": "text",
                "mime_type": "application/json",
                "schema": LESSON_PLAN_SCHEMA
            }
        )
        raw = result.output_text
    else:
        response = client.models.generate_content(
            model=model_name,
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=LESSON_PLAN_SCHEMA
            )
        )
        raw = response.text

    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1]
    if raw.endswith("```"):
        raw = raw[:-3]

    return json.loads(raw.strip()), model_name, available_models

# ---------------------------------------------------------
# WORD EXPORT
# ---------------------------------------------------------
def set_cell_background(cell, fill_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tc_pr.append(shd)

def bullets(items):
    if isinstance(items, list):
        return "\n".join(f"• {i}" for i in items)
    return str(items)

def generate_docx(meta, plan):
    doc = docx.Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.6)
        sec.right_margin = Inches(0.6)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = title_p.add_run("THE ADITYA BIRLA PUBLIC SCHOOL\n")
    run1.bold = True
    run1.font.size = Pt(14)
    run1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    run2 = title_p.add_run(
        "BAIKUNTH, TAH- TILDA, DIST. RAIPUR (C.G.) – 493116\nLesson Plan\n"
    )
    run2.font.size = Pt(10)
    run2.bold = True

    info = doc.add_table(rows=3, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.style = "Table Grid"
    rows_data = [
        [f"Name of the Chapter: {meta['chapter']}", f"Subject: {meta['subject']}"],
        [f"Class: {meta['grade']} {meta['section']}", f"Month: {meta['month']}"],
        [f"Teacher: {meta['teacher']}", f"No. of Periods: {meta['periods']}"],
    ]
    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            cell = info.cell(r_idx, c_idx)
            cell.text = val
            set_cell_background(cell, "F7FAFC")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_section(title, content):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        left = table.cell(0, 0)
        left.width = Inches(2.2)
        left.paragraphs[0].add_run(title).bold = True
        set_cell_background(left, "EDF2F7")

        right = table.cell(0, 1)
        right.width = Inches(4.8)
        if isinstance(content, list):
            right.paragraphs[0].text = ""
            for item in content:
                right.add_paragraph(f"• {item}")
        else:
            right.paragraphs[0].text = str(content)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_section("Curriculum Goal (NCF-SE 2023)", plan.get("curriculum_goal", ""))
    add_section("Relevant Competencies", plan.get("relevant_competencies", []))
    add_section("Learning Objectives", plan.get("learning_objectives", []))
    add_section("Expected Learning Outcomes", plan.get("expected_learning_outcomes", []))

    if plan.get("formulas_and_equations"):
        add_section("Key Formulas, Laws & Equations", plan["formulas_and_equations"])

    add_section("Teaching Methodology", plan.get("teaching_methodology", []))

    aids = ("Teaching Aids:\n" + bullets(plan.get("teaching_aids", []))
            + "\n\nArt Integration:\n" + bullets(plan.get("art_integration", [])))
    add_section("Teaching Aids & Integration of Arts", aids)

    add_section("Connecting Previous Knowledge", plan.get("previous_knowledge", []))
    add_section("Innovative Techniques", plan.get("innovative_techniques", []))

    cp_lines = []
    for item in plan.get("content_points", []):
        cp_lines.append(f"{item.get('section', '')}:")
        for topic in item.get("topics", []):
            cp_lines.append(f"   • {topic}")
    add_section("Content / Teaching Points", "\n".join(cp_lines))

    add_section("Project / Experiential Learning", plan.get("projects_experiential", []))
    add_section("Skills Acquired", plan.get("skills_acquired", []))
    add_section("Values Inculcated", plan.get("values_inculcated", []))

    ass = plan.get("multiple_assessment", {})
    ass_text = (
        "Oral Questions:\n" + bullets(ass.get("oral_questions", []))
        + "\n\nWorksheet:\n" + bullets(ass.get("worksheet", []))
        + "\n\nPractical Assessment:\n" + bullets(ass.get("practical", []))
        + "\n\nExit Ticket:\n" + bullets(ass.get("exit_ticket", []))
    )
    add_section("Multiple / Periodic Assessment", ass_text)

    add_section("Class Work", plan.get("class_work", []))
    add_section("Home Work", plan.get("home_work", []))

    rem = plan.get("remedial_measures", {})
    rem_text = ("Slow Learners:\n" + bullets(rem.get("slow_learners", []))
                + "\n\nAdvanced Learners:\n" + bullets(rem.get("advanced_learners", [])))
    add_section("Remedial Measures", rem_text)

    res = plan.get("resources", {})
    res_text = ("Books:\n" + bullets(res.get("books", []))
                + "\n\nWebsites:\n" + bullets(res.get("websites", []))
                + "\n\nVideos:\n" + bullets(res.get("videos", [])))
    add_section("Resources & References", res_text)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    sig = doc.add_table(rows=1, cols=3)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig.cell(0, 0).text = "___________________\nSubject Teacher"
    sig.cell(0, 1).text = "___________________\nHOD"
    sig.cell(0, 2).text = "___________________\nPrincipal"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# PDF EXPORT
# ---------------------------------------------------------
def esc(text):
    return xml_escape(str(text))

def html_bullets(items):
    if isinstance(items, list):
        return "<br/>".join(f"• {esc(i)}" for i in items)
    return esc(items)

def generate_pdf(meta, plan):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=28, leftMargin=28, topMargin=28, bottomMargin=28,
    )
    story = []

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleStyle", parent=styles["Heading1"], fontName="Helvetica-Bold",
        fontSize=13, alignment=1, textColor=colors.HexColor("#1A365D"),
    )
    sub_style = ParagraphStyle(
        "SubStyle", parent=styles["Normal"], fontName="Helvetica",
        fontSize=8.5, alignment=1,
    )
    body = ParagraphStyle(
        "BodyStyle", parent=styles["Normal"], fontName="Helvetica",
        fontSize=8, leading=11,
    )

    story.append(Paragraph("THE ADITYA BIRLA PUBLIC SCHOOL", title_style))
    story.append(Paragraph(
        "BAIKUNTH, TAH- TILDA, DIST. RAIPUR (C.G.) – 493116<br/><b>Lesson Plan</b>",
        sub_style,
    ))
    story.append(Spacer(1, 8))

    info_data = [
        [Paragraph(f"<b>Chapter:</b> {esc(meta['chapter'])}", body),
         Paragraph(f"<b>Subject:</b> {esc(meta['subject'])}", body)],
        [Paragraph(f"<b>Class:</b> {esc(meta['grade'])} {esc(meta['section'])}", body),
         Paragraph(f"<b>Month:</b> {esc(meta['month'])}", body)],
        [Paragraph(f"<b>Teacher:</b> {esc(meta['teacher'])}", body),
         Paragraph(f"<b>No. of Periods:</b> {esc(meta['periods'])}", body)],
    ]
    info_table = Table(info_data, colWidths=[260, 260])
    info_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7FAFC")),
        ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#CBD5E0")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 6))

    cp_html = ""
    for item in plan.get("content_points", []):
        cp_html += f"<b>{esc(item.get('section', ''))}</b><br/>"
        for topic in item.get("topics", []):
            cp_html += f"• {esc(topic)}<br/>"

    ass = plan.get("multiple_assessment", {})
    ass_html = (
        f"<b>Oral Questions:</b><br/>{html_bullets(ass.get('oral_questions', []))}<br/><br/>"
        f"<b>Worksheet:</b><br/>{html_bullets(ass.get('worksheet', []))}<br/><br/>"
        f"<b>Practical:</b><br/>{html_bullets(ass.get('practical', []))}<br/><br/>"
        f"<b>Exit Ticket:</b><br/>{html_bullets(ass.get('exit_ticket', []))}"
    )

    rem = plan.get("remedial_measures", {})
    rem_html = (
        f"<b>Slow Learners:</b><br/>{html_bullets(rem.get('slow_learners', []))}<br/><br/>"
        f"<b>Advanced Learners:</b><br/>{html_bullets(rem.get('advanced_learners', []))}"
    )

    res = plan.get("resources", {})
    res_html = (
        f"<b>Books:</b><br/>{html_bullets(res.get('books', []))}<br/><br/>"
        f"<b>Websites:</b><br/>{html_bullets(res.get('websites', []))}<br/><br/>"
        f"<b>Videos:</b><br/>{html_bullets(res.get('videos', []))}"
    )

    aids_html = (
        f"<b>Teaching Aids:</b><br/>{html_bullets(plan.get('teaching_aids', []))}<br/><br/>"
        f"<b>Art Integration:</b><br/>{html_bullets(plan.get('art_integration', []))}"
    )

    rows = [
        ["Curriculum Goal (NCF-SE 2023)", esc(plan.get("curriculum_goal", ""))],
        ["Relevant Competencies", html_bullets(plan.get("relevant_competencies", []))],
        ["Learning Objectives", html_bullets(plan.get("learning_objectives", []))],
        ["Expected Learning Outcomes", html_bullets(plan.get("expected_learning_outcomes", []))],
    ]
    if plan.get("formulas_and_equations"):
        rows.append(["Key Formulas & Equations", html_bullets(plan["formulas_and_equations"])])
    rows.extend([
        ["Teaching Methodology", html_bullets(plan.get("teaching_methodology", []))],
        ["Teaching Aids & Art Integration", aids_html],
        ["Connecting Previous Knowledge", html_bullets(plan.get("previous_knowledge", []))],
        ["Innovative Techniques", html_bullets(plan.get("innovative_techniques", []))],
        ["Content / Teaching Points", cp_html],
        ["Project / Experiential Learning", html_bullets(plan.get("projects_experiential", []))],
        ["Skills Acquired", html_bullets(plan.get("skills_acquired", []))],
        ["Values Inculcated", html_bullets(plan.get("values_inculcated", []))],
        ["Multiple / Periodic Assessment", ass_html],
        ["Class Work", html_bullets(plan.get("class_work", []))],
        ["Home Work", html_bullets(plan.get("home_work", []))],
        ["Remedial Measures", rem_html],
        ["Resources & References", res_html],
    ])

    details_data = [
        [Paragraph(f"<b>{label}</b>", body), Paragraph(value, body)]
        for label, value in rows
    ]

    details_table = Table(details_data, colWidths=[135, 385], repeatRows=0)
    details_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E0")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EDF2F7")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(details_table)
    story.append(Spacer(1, 14))

    sig_table = Table(
        [["___________________\nSubject Teacher",
          "___________________\nHOD",
          "___________________\nPrincipal"]],
        colWidths=[173, 173, 173],
    )
    sig_table.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    story.append(sig_table)

    doc.build(story)
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# GENERATION EXECUTION
# ---------------------------------------------------------
st.divider()
if st.button("✨ Generate AI Lesson Plan & Mind Map", type="primary", use_container_width=True):
    if not active_api_key:
        st.error("⚠️ Gemini API key is missing.")
    elif uploaded_pdf is not None and uploaded_pdf.size > 25 * 1024 * 1024:
        st.error("⚠️ File exceeds 25 MB.")
    else:
        chapter_content = ""

        if uploaded_pdf is not None:
            with st.spinner("Extracting text from PDF..."):
                chapter_content = extract_text_cached(uploaded_pdf.getvalue())
            if not chapter_content.strip():
                st.warning("No selectable text found in the PDF. Please paste chapter text manually.")
        elif pasted_text.strip():
            chapter_content = pasted_text.strip()
        else:
            chapter_content = f"NCERT content for {subject}, class {grade}, chapter {chapter}"

        with st.spinner("Generating lesson plan..."):
            try:
                plan_data, used_model, available_models = generate_ai_lesson_plan(
                    active_api_key, subject, grade, section, chapter, month, periods, chapter_content
                )

                st.session_state["plan_data"] = plan_data
                st.session_state["meta"] = {
                    "teacher": teacher_name,
                    "subject": subject,
                    "grade": grade,
                    "section": section,
                    "chapter": chapter,
                    "periods": periods,
                    "month": month
                }

                st.success(f"✅ Lesson plan generated using {used_model}")
                with st.expander("Available models detected for this API key"):
                    st.code("\n".join(available_models[:50]))
            except Exception as e:
                st.error("Lesson plan generation failed.")
                st.code(str(e))

# ---------------------------------------------------------
# OUTPUT PRESENTATION
# ---------------------------------------------------------
if "plan_data" in st.session_state:
    data = st.session_state["plan_data"]
    meta = st.session_state["meta"]

    st.subheader(f"📋 Preview: {meta['chapter']} ({meta['subject']})")
    st.markdown(f"**Curriculum Goal:** {data.get('curriculum_goal', '')}")

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("### Learning Objectives")
        for obj in data.get("learning_objectives", []):
            st.markdown(f"- {obj}")
    with col2:
        st.markdown("### Expected Outcomes")
        for outcome in data.get("expected_learning_outcomes", []):
            st.markdown(f"- {outcome}")

    if data.get("formulas_and_equations"):
        st.divider()
        st.subheader("📐 Key Formulas & Laws Extracted")
        for formula in data["formulas_and_equations"]:
            st.info(f"⚡ {formula}")

    st.divider()
    st.subheader(f"🧠 Visual Mind Map: {meta['chapter']}")
    try:
        dot = graphviz.Digraph(comment=meta["chapter"])
        dot.attr(rankdir="LR", splines="ortho")
        dot.attr("node", style="filled", fillcolor="#EBF8FF",
                 color="#2B6CB0", fontname="Helvetica", fontsize="10", shape="box")
        dot.node("CENTER", meta["chapter"], fillcolor="#2B6CB0", fontcolor="white")

        for s_idx, item in enumerate(data.get("content_points", [])):
            section_id = f"SEC{s_idx}"
            dot.node(section_id, item.get("section", "Core Concepts"),
                     shape="ellipse", fillcolor="#E2E8F0")
            dot.edge("CENTER", section_id)
            for t_idx, topic in enumerate(item.get("topics", [])):
                topic_id = f"SEC{s_idx}_T{t_idx}"
                dot.node(topic_id, topic, fillcolor="#FFFFFF")
                dot.edge(section_id, topic_id)

        st.graphviz_chart(dot, use_container_width=True)
    except Exception:
        st.warning("Mind map rendering skipped. Downloads below contain all details.")

    st.divider()
    safe_chapter = meta["chapter"].replace(" ", "_").replace("/", "-")

    col_pdf, col_docx = st.columns(2)
    with col_pdf:
        st.download_button(
            label="📄 Download Official PDF",
            data=generate_pdf(meta, data),
            file_name=f"ABPS_Lesson_Plan_{meta['grade']}_{safe_chapter}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    with col_docx:
        st.download_button(
            label="📝 Download Editable Word (.docx)",
            data=generate_docx(meta, data),
            file_name=f"ABPS_Lesson_Plan_{meta['grade']}_{safe_chapter}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True,
        )
